#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import json
import sys
import os
from datetime import datetime

# Logging per debug
def debug_log(message):
    print(f"DEBUG: {message}", file=sys.stderr)

debug_log("Script Python avviato")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from docx.dml.color import ColorFormat
    from docx.enum.dml import MSO_THEME_COLOR_INDEX
    debug_log("Importazioni docx completate")
except ImportError as e:
    debug_log(f"Errore importazione docx: {e}")
    print(f"ERROR: python-docx non installato: {e}")
    sys.exit(1)

def add_table_border(table):
    """Aggiunge bordi alla tabella"""
    try:
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # Crea elemento bordi
        tblBorders = OxmlElement('w:tblBorders')
        
        # Definisci i bordi
        borders = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
        for border in borders:
            border_el = OxmlElement(f'w:{border}')
            border_el.set(qn('w:val'), 'single')
            border_el.set(qn('w:sz'), '4')
            border_el.set(qn('w:color'), '000000')
            tblBorders.append(border_el)
        
        tblPr.append(tblBorders)
        debug_log("Bordi tabella aggiunti")
    except Exception as e:
        debug_log(f"Errore aggiunta bordi: {e}")

def format_cell_text(cell, text, bold_if_total=False):
    """Formatta il testo base di una cella"""
    try:
        # Pulisci la cella
        cell.text = ''
        
        # Aggiungi paragrafo
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.space_before = Pt(2)
        
        # Aggiungi il testo
        run = paragraph.add_run(str(text))
        run.font.size = Pt(8)
        run.font.name = 'Arial'
        
        # Grassetto se è la riga totale
        if bold_if_total and '**' in str(text):
            run.font.bold = True
        
    except Exception as e:
        debug_log(f"Errore formattazione cella: {e}")
        cell.text = str(text)

def format_multi_line_cell(cell, text, separator):
    """Formatta una cella con testo multi-linea"""
    try:
        if not text or text in ['N/A', '']:
            format_cell_text(cell, text)
            return
        
        # Pulisci la cella
        cell.text = ''
        
        # Prepara il testo per la divisione
        text_str = str(text)
        
        # Gestisci diversi tipi di separatori
        if separator == '•':
            # Per caratteristiche HP - NUOVO: gestisce anche i trattini
            if '•' in text_str:
                # Prima dividi per •
                lines = text_str.split('•')
                lines = [line.strip() for line in lines if line.strip()]
                
                # Poi per ogni linea, controlla se ha trattini e subdividi
                final_lines = []
                for line in lines:
                    if ' - ' in line:
                        # Dividi per trattini mantenendo il bullet
                        sub_lines = line.split(' - ')
                        for i, sub_line in enumerate(sub_lines):
                            if sub_line.strip():
                                if i == 0:
                                    final_lines.append(sub_line.strip())
                                else:
                                    final_lines.append(f"- {sub_line.strip()}")
                    else:
                        final_lines.append(line)
                lines = final_lines
            elif ' - ' in text_str:
                # Se non ci sono •, ma ci sono trattini, dividi per trattini
                lines = text_str.split(' - ')
                processed_lines = []
                for i, line in enumerate(lines):
                    if line.strip():
                        if i == 0:
                            processed_lines.append(line.strip())
                        else:
                            processed_lines.append(f"- {line.strip()}")
                lines = processed_lines
            else:
                lines = [text_str]
        elif separator == ';':
            # Per motivi HP
            if ';' in text_str:
                lines = text_str.split(';')
                lines = [line.strip() for line in lines if line.strip()]
            else:
                lines = [text_str]
        elif separator == ',':
            # Per codici e categoria pericolo
            if ',' in text_str:
                lines = text_str.split(',')
                lines = [line.strip() for line in lines if line.strip()]
            elif '\n' in text_str:
                lines = text_str.split('\n')
                lines = [line.strip() for line in lines if line.strip()]
            else:
                lines = [text_str]
        else:
            lines = [text_str]
        
        # Aggiungi le linee alla cella
        for i, line in enumerate(lines):
            if i == 0:
                # Prima linea - usa il paragrafo esistente
                paragraph = cell.paragraphs[0]
            else:
                # Linee successive - aggiungi nuovi paragrafi
                paragraph = cell.add_paragraph()
            
            # Formatta il paragrafo
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.line_spacing = 1.0
            
            # Aggiungi il testo
            if separator == '•':
                # Per caratteristiche HP - gestione migliorata
                if i == 0 and not line.startswith('•') and not line.startswith('-'):
                    # Prima linea senza prefisso
                    run = paragraph.add_run(line)
                elif line.startswith('- '):
                    # Linea che inizia già con trattino
                    run = paragraph.add_run(line)
                elif i > 0 and not line.startswith('•') and not line.startswith('-'):
                    # Linee successive senza prefisso, aggiungi bullet
                    run = paragraph.add_run(f"• {line}")
                else:
                    # Mantieni il formato originale
                    run = paragraph.add_run(line)
            elif separator == ';' and len(lines) > 1:
                # Per motivi HP, se ci sono più linee
                if i == len(lines) - 1:
                    # Ultima linea senza punto e virgola
                    run = paragraph.add_run(line)
                else:
                    # Linee intermedie con punto e virgola
                    run = paragraph.add_run(f"{line};")
            elif separator == ',' and len(lines) > 1:
                # Per codici e categoria pericolo
                run = paragraph.add_run(line)
            else:
                run = paragraph.add_run(line)
            
            # Formatta il run
            run.font.size = Pt(8)
            run.font.name = 'Arial'
            
            # Grassetto se è la riga totale
            if '**' in line:
                run.font.bold = True
                
            # Colore diverso per le categorie di pericolo
            if separator == ',' and ' - ' in line:
                # Se è un codice con categoria (es. "Acute Tox. 4 - H302")
                run.font.color.rgb = RGBColor(0, 100, 0)  # Verde scuro per le categorie
        
    except Exception as e:
        debug_log(f"Errore formattazione multi-linea: {e}")
        format_cell_text(cell, text)

def add_complete_header_content(doc, metadata):
    """Aggiunge l'intestazione completa nel corpo del documento (non nell'header di Word)"""
    try:
        # Trova il logo nella cartella del software
        logo_path = None
        script_dir = os.path.dirname(os.path.abspath(__file__))
        possible_logo_paths = [
            os.path.join(script_dir, 'logo_sea.png'),
            os.path.join(script_dir, '..', 'logo_sea.png'),
            os.path.join(script_dir, 'assets', 'logo_sea.png'),
            os.path.join(script_dir, '..', 'assets', 'logo_sea.png')
        ]
        
        for path in possible_logo_paths:
            if os.path.exists(path):
                logo_path = path
                break
        
        # Estrai informazioni dai metadata
        rapporto_prova = 'N/A'
        codice_eer = 'N/A'
        
        if isinstance(metadata.get('infoCertificato'), dict):
            rapporto_prova = metadata['infoCertificato'].get('rapportoProva', 'N/A')
            codice_eer = metadata['infoCertificato'].get('codiceEER', 'N/A')
        
        # 1. LOGO E INFORMAZIONI AZIENDALI
        header_table = doc.add_table(rows=1, cols=2)
        header_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Cella logo (sinistra)
        logo_cell = header_table.cell(0, 0)
        logo_para = logo_cell.paragraphs[0]
        logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_para.paragraph_format.space_after = Pt(0)
        
        # Aggiungi logo se disponibile
        if logo_path:
            try:
                run = logo_para.add_run()
                run.add_picture(logo_path, width=Inches(1.0))
                debug_log(f"Logo aggiunto: {logo_path}")
            except Exception as e:
                debug_log(f"Errore aggiunta logo: {e}")
                run = logo_para.add_run("SEA SRLS")
                run.font.size = Pt(8)
                run.font.name = 'Arial'
        else:
            run = logo_para.add_run("SEA SRLS")
            run.font.size = Pt(8)
            run.font.name = 'Arial'
        
        # Cella informazioni (destra del logo)
        info_cell = header_table.cell(0, 1)
        info_para = info_cell.paragraphs[0]
        info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        info_para.paragraph_format.space_after = Pt(0)
        
        # Informazioni aziendali
        company_info = [
            "SEA SRLS",
            "Via Madre Teresa di Calcutta, 100",
            "62029 Tolentino (MC)", 
            "Tel.: 0733/974947    Fax: 0733/967097",
            "e-mail: info@seatoientino.it",
            "internet: www.seatoientino.it"
        ]
        
        for i, line in enumerate(company_info):
            if i > 0:
                info_para.add_run("\n")
            run = info_para.add_run(line)
            run.font.size = Pt(7)
            run.font.name = 'Arial'
        
        # Rimuovi bordi dalla tabella header
        header_table._tbl.tblPr.append(parse_xml(r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>'))
        
        # 2. SPAZIO
        space_para = doc.add_paragraph()
        space_para.paragraph_format.space_after = Pt(2)
        
        # 3. ALLEGATO AL RAPPORTO DI PROVA
        allegato_p = doc.add_paragraph()
        allegato_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        allegato_p.paragraph_format.space_after = Pt(1)
        allegato_run = allegato_p.add_run(f"Allegato al Rapporto di Prova N° {rapporto_prova}")
        allegato_run.font.size = Pt(9)
        allegato_run.font.name = 'Arial'
        
        # 4. TITOLO PRINCIPALE - NON IN GRASSETTO
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_after = Pt(1)
        title_run = title_p.add_run("PARERI ED INTERPRETAZIONI – NON OGGETTO DELL'ACCREDITAMENTO ACCREDIA")
        title_run.font.size = Pt(10)
        title_run.font.name = 'Arial'
        
        # 5. SOTTOTITOLO - NON IN GRASSETTO
        subtitle_p = doc.add_paragraph()
        subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_p.paragraph_format.space_after = Pt(1)
        subtitle_run = subtitle_p.add_run(f"CARATTERIZZAZIONE DEL RIFIUTO CODICE EER {codice_eer}")
        subtitle_run.font.size = Pt(9)
        subtitle_run.font.name = 'Arial'
        
        # 6. TESTO DI RIFERIMENTO NORMATIVO
        ref_text = f"""(redatto ai sensi del Decreto Legislativo 152/2006 e s.m.i., della Linee Guida SNPA 2021 Quadro 2.2 - Delibera n°105/2021
Approvato con D.D. n°LG/2021-53 del 13/12/2021, del Regolamento (UE) 1357/2014, Regolamento (UE) 2017/997))"""
        
        ref_p = doc.add_paragraph()
        ref_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ref_p.paragraph_format.space_after = Pt(2)
        ref_run = ref_p.add_run(ref_text)
        ref_run.font.size = Pt(7)
        ref_run.font.name = 'Arial'
        
        # 7. LINEA DI SEPARAZIONE
        separator_p = doc.add_paragraph()
        separator_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        separator_p.paragraph_format.space_after = Pt(2)
        separator_run = separator_p.add_run("_" * 80)
        separator_run.font.size = Pt(7)
        separator_run.font.name = 'Arial'
        
        # SPAZIO AGGIUNTIVO DOPO SEPARATORE
        space_after_separator = doc.add_paragraph()
        space_after_separator.paragraph_format.space_after = Pt(3)
        
        debug_log("Header completo aggiunto nel corpo del documento")
        
    except Exception as e:
        debug_log(f"Errore creazione header completo: {e}")

def create_classification_tables(doc, metadata, risultati):
    """Crea le tabelle di classificazione come nel formato SEA (senza bordi veri)"""
    try:
        # TABELLA 1: Dati relativi al Campione e al Rapporto di Prova
        table1_title = doc.add_paragraph()
        table1_title.paragraph_format.space_before = Pt(1)
        table1_title.paragraph_format.space_after = Pt(1)
        table1_title_run = table1_title.add_run("Tabella 1. Dati relativi al Campione e al Rapporto di Prova")
        table1_title_run.font.size = Pt(8)
        table1_title_run.font.bold = True
        table1_title_run.font.name = 'Arial'
        
        # Estrai dati dai metadata
        rapporto_prova = 'N/A'
        numero_campionamento = 'N/A'
        data_campionamento = 'N/A'
        committente = 'Non specificato'
        descrizione = 'Prodotti ad H2O non conformi o obsoleti non recuperabili'
        codice_eer = 'N/A'
        
        if isinstance(metadata.get('infoCertificato'), dict):
            rapporto_prova = metadata['infoCertificato'].get('rapportoProva', 'N/A')
            numero_campionamento = metadata['infoCertificato'].get('numeroCampionamento', 'N/A')
            data_campionamento = metadata['infoCertificato'].get('dataCampionamento', 'N/A')
            descrizione = metadata['infoCertificato'].get('descrizione', descrizione)
            codice_eer = metadata['infoCertificato'].get('codiceEER', 'N/A')
        
        committente = metadata.get('committente', 'Non specificato')
        
        # Invece di tabella, usa un approccio diretto per ridurre spazi
        desc_p = doc.add_paragraph(f"Descrizione Prodotti: {descrizione} EER {codice_eer}")
        desc_p.paragraph_format.space_after = Pt(0)
        desc_p.paragraph_format.space_before = Pt(0)
        for run in desc_p.runs:
            run.font.size = Pt(8)
            run.font.name = 'Arial'
        
        ref_p = doc.add_paragraph(f"Riferimento del Verbale di Campionamento: {numero_campionamento}")
        ref_p.paragraph_format.space_after = Pt(0)
        ref_p.paragraph_format.space_before = Pt(0)
        for run in ref_p.runs:
            run.font.size = Pt(8)
            run.font.name = 'Arial'
        
        data_p = doc.add_paragraph(f"Data di campionamento: {data_campionamento}")
        data_p.paragraph_format.space_after = Pt(0)
        data_p.paragraph_format.space_before = Pt(0)
        for run in data_p.runs:
            run.font.size = Pt(8)
            run.font.name = 'Arial'
        
        comm_p = doc.add_paragraph(f"Committente: {committente}")
        comm_p.paragraph_format.space_after = Pt(3)
        comm_p.paragraph_format.space_before = Pt(0)
        for run in comm_p.runs:
            run.font.size = Pt(8)
            run.font.name = 'Arial'
        
        # TABELLA 2: Descrizione del Rifiuto
        table2_title = doc.add_paragraph()
        table2_title.paragraph_format.space_before = Pt(1)
        table2_title.paragraph_format.space_after = Pt(1)
        table2_title_run = table2_title.add_run("Tabella 2. Descrizione del Rifiuto")
        table2_title_run.font.size = Pt(8)
        table2_title_run.font.bold = True
        table2_title_run.font.name = 'Arial'
        
        # Applica anche spaziatura zero agli altri paragrafi delle tabelle - FONT 8PT
        desc_prod_p = doc.add_paragraph()
        desc_prod_p.paragraph_format.space_after = Pt(0)
        desc_prod_p.paragraph_format.space_before = Pt(0)
        desc_prod_p.paragraph_format.line_spacing = 0.8
        desc_prod_run = desc_prod_p.add_run(f"Descrizione del ciclo produttivo: Rifiuto derivante dalla produzione fuori specifica")
        desc_prod_run.font.size = Pt(8)
        desc_prod_run.font.name = 'Arial'
        
        eer_p = doc.add_paragraph()
        eer_p.paragraph_format.space_after = Pt(0)
        eer_p.paragraph_format.space_before = Pt(0)
        eer_p.paragraph_format.line_spacing = 0.8
        eer_run = eer_p.add_run(f"EER: {codice_eer} {descrizione}")
        eer_run.font.size = Pt(8)
        eer_run.font.name = 'Arial'
        
        # Matrice presa dai metadata - FONT 8PT
        matrice = metadata.get('infoCertificato', {}).get('matrice', 'Gestione rifiuto')
        matrice_p = doc.add_paragraph()
        matrice_p.paragraph_format.space_after = Pt(3)
        matrice_p.paragraph_format.space_before = Pt(0)
        matrice_p.paragraph_format.line_spacing = 0.8
        matrice_run = matrice_p.add_run(f"Matrice: {matrice}")
        matrice_run.font.size = Pt(8)
        matrice_run.font.name = 'Arial'
        
        # TABELLA 3: Aspetto e caratteristiche chimico-fisiche del campione
        table3_title = doc.add_paragraph()
        table3_title.paragraph_format.space_before = Pt(1)
        table3_title.paragraph_format.space_after = Pt(1)
        table3_title_run = table3_title.add_run("Tabella 3. Aspetto e caratteristiche chimico-fisiche del campione")
        table3_title_run.font.size = Pt(8)
        table3_title_run.font.bold = True
        table3_title_run.font.name = 'Arial'
        
        # Estrai caratteristiche fisiche dai metadata se disponibili
        caratteristiche_fisiche = metadata.get('caratteristicheFisiche', {})
        
        # Applica spaziatura zero anche alle caratteristiche fisiche - FONT 8PT
        colore_p = doc.add_paragraph()
        colore_p.paragraph_format.space_after = Pt(0)
        colore_p.paragraph_format.space_before = Pt(0)
        colore_p.paragraph_format.line_spacing = 0.8
        colore_run = colore_p.add_run(f"Colore: {caratteristiche_fisiche.get('colore', 'bianco')}")
        colore_run.font.size = Pt(8)
        colore_run.font.name = 'Arial'
        
        odore_p = doc.add_paragraph()
        odore_p.paragraph_format.space_after = Pt(0)
        odore_p.paragraph_format.space_before = Pt(0)
        odore_p.paragraph_format.line_spacing = 0.8
        odore_run = odore_p.add_run(f"Odore: {caratteristiche_fisiche.get('odore', 'chimico non definibile')}")
        odore_run.font.size = Pt(8)
        odore_run.font.name = 'Arial'
        
        stato_p = doc.add_paragraph()
        stato_p.paragraph_format.space_after = Pt(0)
        stato_p.paragraph_format.space_before = Pt(0)
        stato_p.paragraph_format.line_spacing = 0.8
        stato_run = stato_p.add_run(f"Stato fisico: {caratteristiche_fisiche.get('statoFisico', 'fangoso palpabile')}")
        stato_run.font.size = Pt(8)
        stato_run.font.name = 'Arial'
        
        ph_p = doc.add_paragraph()
        ph_p.paragraph_format.space_after = Pt(0)
        ph_p.paragraph_format.space_before = Pt(0)
        ph_p.paragraph_format.line_spacing = 0.8
        ph_run = ph_p.add_run(f"pH: {caratteristiche_fisiche.get('ph', '5,0')}")
        ph_run.font.size = Pt(8)
        ph_run.font.name = 'Arial'
        
        res105_p = doc.add_paragraph()
        res105_p.paragraph_format.space_after = Pt(0)
        res105_p.paragraph_format.space_before = Pt(0)
        res105_p.paragraph_format.line_spacing = 0.8
        res105_run = res105_p.add_run(f"Residuo a 105 °C: {caratteristiche_fisiche.get('residuo105', '49,1')}")
        res105_run.font.size = Pt(8)
        res105_run.font.name = 'Arial'
        
        res180_p = doc.add_paragraph()
        res180_p.paragraph_format.space_after = Pt(0)
        res180_p.paragraph_format.space_before = Pt(0)
        res180_p.paragraph_format.line_spacing = 0.8
        res180_run = res180_p.add_run(f"Residuo 180°C: {caratteristiche_fisiche.get('residuo180', '46,8')}")
        res180_run.font.size = Pt(8)
        res180_run.font.name = 'Arial'
        
        res600_p = doc.add_paragraph()
        res600_p.paragraph_format.space_after = Pt(3)
        res600_p.paragraph_format.space_before = Pt(0)
        res600_p.paragraph_format.line_spacing = 0.8
        res600_run = res600_p.add_run(f"Residuo 600 °C: {caratteristiche_fisiche.get('residuo600', '0,3')}")
        res600_run.font.size = Pt(8)
        res600_run.font.name = 'Arial'
        
        debug_log("Tabelle di classificazione create (formato paragrafi)")
        
    except Exception as e:
        debug_log(f"Errore creazione tabelle classificazione: {e}")

def add_footer_with_page_numbers(doc, metadata):
    """Aggiunge footer con data rapporto a sinistra e numero pagina a destra"""
    try:
        # Estrai data rapporto dai metadata
        data_rapporto = 'N/A'
        if isinstance(metadata.get('infoCertificato'), dict):
            data_rapporto = metadata['infoCertificato'].get('dataCampionamento', 'N/A')
        
        # Accedi al footer del documento
        section = doc.sections[0]
        footer = section.footer
        
        # Pulisci il footer esistente
        footer._element.clear()
        
        # Crea un nuovo paragrafo nel footer
        footer_para = footer.add_paragraph()
        
        # Imposta tab stop per allineamento
        from docx.shared import Inches
        from docx.enum.text import WD_TAB_ALIGNMENT
        tab_stops = footer_para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        
        # Aggiungi il testo della data a sinistra
        data_run = footer_para.add_run(f"Data Rap. Prova: {data_rapporto}")
        data_run.font.size = Pt(9)
        data_run.font.name = 'Arial'
        data_run.font.color.rgb = None  # Assicura che sia nero
        
        # Aggiungi tab per andare a destra
        footer_para.add_run("\t")
        
        # Aggiungi testo "Pagina" e numero pagina
        page_text_run = footer_para.add_run("Pagina ")
        page_text_run.font.size = Pt(9)
        page_text_run.font.name = 'Arial'
        page_text_run.font.color.rgb = None
        
        # Aggiungi campo numero pagina corrente
        page_num_run = footer_para.add_run()
        page_num_run.font.size = Pt(9)
        page_num_run.font.name = 'Arial'
        page_num_run.font.color.rgb = None
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        page_num_run._r.append(fldChar1)
        page_num_run._r.append(instrText)
        page_num_run._r.append(fldChar2)
        
        # Aggiungi " di " e numero totale pagine
        di_run = footer_para.add_run(" di ")
        di_run.font.size = Pt(9)
        di_run.font.name = 'Arial'
        di_run.font.color.rgb = None
        
        # Aggiungi campo numero totale pagine
        total_pages_run = footer_para.add_run()
        total_pages_run.font.size = Pt(9)
        total_pages_run.font.name = 'Arial'
        total_pages_run.font.color.rgb = None
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        
        instrText2 = OxmlElement('w:instrText')
        instrText2.text = "NUMPAGES"
        
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        
        total_pages_run._r.append(fldChar3)
        total_pages_run._r.append(instrText2)
        total_pages_run._r.append(fldChar4)
        
        debug_log("Footer con data rapporto e paginazione completa aggiunto")
        
    except Exception as e:
        debug_log(f"Errore creazione footer: {e}")

def add_page_break_with_header(doc, metadata):
    """Aggiunge un'interruzione di pagina e ripete l'header MANUALMENTE"""
    try:
        # 1. Aggiungi interruzione di pagina
        doc.add_page_break()
        
        # 2. Ripeti TUTTO l'header completo
        add_complete_header_content(doc, metadata)
        
        debug_log("Page break con header completo aggiunto")
        
    except Exception as e:
        debug_log(f"Errore page break con header: {e}")

def add_document_title_section(doc, metadata):
    """Aggiunge la sezione del titolo del documento nel corpo (dopo header)"""
    try:
        # Estrai informazioni dai metadata
        rapporto_prova = 'N/A'
        data_campionamento = 'N/A'
        
        if isinstance(metadata.get('infoCertificato'), dict):
            rapporto_prova = metadata['infoCertificato'].get('rapportoProva', 'N/A')
            data_campionamento = metadata['infoCertificato'].get('dataCampionamento', 'N/A')
        
        # Titolo sezione classificazione - grassetto solo questo
        class_title_p = doc.add_paragraph()
        class_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        class_title_p.paragraph_format.space_after = Pt(4)
        class_title_run = class_title_p.add_run("CLASSIFICAZIONE E/O CARATTERIZZAZIONE DEL RIFIUTO")
        class_title_run.font.size = Pt(9)
        class_title_run.font.bold = True
        class_title_run.font.name = 'Arial'
        
        # Riferimento al rapporto
        rapport_p = doc.add_paragraph()
        rapport_p.paragraph_format.space_after = Pt(4)
        rapport_run = rapport_p.add_run(f"In relazione al Rapporto di prova n. {rapporto_prova} del {data_campionamento}")
        rapport_run.font.size = Pt(8)
        rapport_run.font.name = 'Arial'
        
        # Testo descrittivo
        desc_text = """Il profilo analitico, determinante per la caratterizzazione è stato scelto con il Produttore, sulla base delle informazioni fornite dallo stesso, inerenti processi di lavoro nonché su ipotesi generalmente ben fondate da parte del produttore di tali rifiuti, ed eventuali schede di sicurezza dei prodotti da cui deriva.

La presente valutazione si riferisce esclusivamente al campione analizzato, ai test eseguiti e ai parametri analizzati.

Per le informazioni (se disponibili) richieste ai punti 1-2-3-4-9-10 nel Riquadro 2.2 delle Linee Guida SNPA approvate con Decreto Direttoriale 47/2021 si rimanda al Rapporto di Prova di cui la presente valutazione è un allegato."""
        
        desc_p = doc.add_paragraph()
        desc_p.paragraph_format.space_after = Pt(4)
        desc_run = desc_p.add_run(desc_text)
        desc_run.font.size = Pt(8)
        desc_run.font.name = 'Arial'
        
        # Nota sulla classificazione
        note_p = doc.add_paragraph()
        note_p.paragraph_format.space_after = Pt(6)
        note_run = note_p.add_run("La classificazione delle sostanze pericolose presenti in miscela, ove non espressamente dichiarato, si intende riferita nell'elenco armonizzato delle sostanze presenti nell'allegato VI del Regolamento CLP.")
        note_run.font.size = Pt(7)
        note_run.font.italic = True
        note_run.font.name = 'Arial'
        
        debug_log("Sezione titolo documento aggiunta nel corpo")
        
    except Exception as e:
        debug_log(f"Errore creazione sezione titolo: {e}")

def create_results_table(doc, risultati):
    """Crea la tabella dei risultati della classificazione (Tabella 4) con struttura CORRETTA per l'anteprima"""
    try:
        debug_log("Creazione tabella risultati")
        
        # Verifica che ci siano risultati
        if not risultati or len(risultati) == 0:
            no_data_p = doc.add_paragraph()
            no_data_run = no_data_p.add_run("Nessun dato di classificazione disponibile")
            no_data_run.font.size = Pt(7)
            no_data_run.font.name = 'Arial'
            return
        
        # Titolo sezione risultati
        results_title = doc.add_paragraph()
        results_title.paragraph_format.space_before = Pt(2)
        results_title.paragraph_format.space_after = Pt(2)
        results_title_run = results_title.add_run("Tabella 4. Risultati della Classificazione")
        results_title_run.font.size = Pt(8)
        results_title_run.font.bold = True
        results_title_run.font.name = 'Arial'
        
        # CORREZIONE: Usa le stesse colonne dell'anteprima JavaScript
        # Le colonne corrette sono: Sostanza, Conc. (ppm), Conc. (%), Caratteristiche HP, CAS, Motivo HP, Codici e Categoria Pericolo
        headers = ['Sostanza', 'Conc. (ppm)', 'Conc. (%)', 'Caratteristiche HP', 'CAS', 'Motivo HP', 'Codici e Categoria Pericolo']
        debug_log(f"Headers utilizzati: {headers}")
        
        # Crea tabella con intestazioni
        results_table = doc.add_table(rows=1, cols=len(headers))
        results_table.style = 'Table Grid'
        
        # Migliora la formattazione della tabella - LARGHEZZA MASSIMA
        results_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Imposta margini minimi per la sezione
        try:
            section = doc.sections[0]
            section.left_margin = Inches(0.5)   # Margine sinistro minimo
            section.right_margin = Inches(0.5)  # Margine destro minimo
            section.top_margin = Inches(0.5)    # Margine superiore minimo
            section.bottom_margin = Inches(0.5) # Margine inferiore minimo
        except Exception as e:
            debug_log(f"Errore impostazione margini: {e}")
        
        # Imposta larghezza colonne per sfruttare tutto lo spazio disponibile
        # Larghezza totale disponibile: circa 7.5 inches (con margini 0.5")
        try:
            col_widths = [
                Inches(1.8),  # Sostanza - più larga
                Inches(0.7),  # Conc. (ppm) - compatta
                Inches(0.7),  # Conc. (%) - compatta
                Inches(2.2),  # Caratteristiche HP - molto più larga
                Inches(0.8),  # CAS - compatta
                Inches(2.5),  # Motivo HP - molto larga
                Inches(2.0)   # Codici e Categoria Pericolo - larga
            ]
            
            for i, width in enumerate(col_widths):
                if i < len(results_table.columns):
                    results_table.columns[i].width = width
        except Exception as e:
            debug_log(f"Errore impostazione larghezza colonne: {e}")
        
        add_table_border(results_table)
        
        # Aggiungi header con formattazione migliorata
        header_cells = results_table.rows[0].cells
        for i, header in enumerate(headers):
            # Pulisci la cella
            header_cells[i].text = ''
            
            # Aggiungi paragrafo con formattazione
            paragraph = header_cells[i].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.space_before = Pt(3)
            
            # Aggiungi il testo dell'header
            run = paragraph.add_run(header)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(255, 255, 255)  # Testo bianco
            
            # Colore di sfondo per gli header
            try:
                cell_shading = header_cells[i]._element.xpath('.//w:shd')[0] if header_cells[i]._element.xpath('.//w:shd') else None
                if not cell_shading:
                    shading_elm = parse_xml(r'<w:shd {} w:fill="366092"/>'.format(nsdecls('w')))
                    header_cells[i]._element.get_or_add_tcPr().append(shading_elm)
            except Exception as e:
                debug_log(f"Errore colore header: {e}")
        
        # Imposta altezza minima per la riga header
        try:
            results_table.rows[0].height = Inches(0.5)
        except Exception as e:
            debug_log(f"Errore altezza header: {e}")
        
        # Aggiungi dati con formattazione avanzata
        for result in risultati:
            row_cells = results_table.add_row().cells
            
            # === SOSTANZA ===
            format_cell_text(row_cells[0], str(result.get('Sostanza', '')), bold_if_total=True)
            
            # === CONC. (PPM) ===
            conc_ppm = result.get('Conc. (ppm)', result.get('Concentrazione (ppm)', ''))
            if isinstance(conc_ppm, (int, float)):
                format_cell_text(row_cells[1], f"{float(conc_ppm):.3f}")
            else:
                format_cell_text(row_cells[1], str(conc_ppm) if conc_ppm else '')
            
            # === CONC. (%) ===
            conc_perc = result.get('Conc. (%)', result.get('Concentrazione (%)', ''))
            if isinstance(conc_perc, (int, float)):
                format_cell_text(row_cells[2], f"{float(conc_perc):.3f}")
            else:
                format_cell_text(row_cells[2], str(conc_perc) if conc_perc else '')
            
            # === CARATTERISTICHE HP ===
            caratteristiche_hp = result.get('Caratteristiche HP', '')
            format_multi_line_cell(row_cells[3], caratteristiche_hp, separator='•')
            
            # === CAS ===
            format_cell_text(row_cells[4], str(result.get('CAS', '')))
            
            # === MOTIVO HP ===
            motivo_hp = result.get('Motivo HP', '')
            format_multi_line_cell(row_cells[5], motivo_hp, separator=';')
            
            # === CODICI E CATEGORIA PERICOLO ===
            codici_categoria = result.get('Codici e Categoria Pericolo', result.get('Frasi H', ''))
            
            # Rimuovi HTML e formatta
            if isinstance(codici_categoria, str):
                codici_categoria = codici_categoria.replace('<br>', '\n').replace('&gt;', '>').replace('&lt;', '<')
            
            format_multi_line_cell(row_cells[6], codici_categoria, separator=',')
            
            # Applica altezza minima alla riga per migliorare la leggibilità
            try:
                row = results_table.rows[len(results_table.rows) - 1]
                row.height = Inches(0.4)  # Altezza minima riga
            except Exception as e:
                debug_log(f"Errore impostazione altezza riga: {e}")
        
        doc.add_paragraph()
        debug_log("Tabella risultati creata con successo")
        
    except Exception as e:
        debug_log(f"Errore creazione tabella risultati: {e}")
        # Aggiungi messaggio di errore nel documento
        error_p = doc.add_paragraph()
        error_run = error_p.add_run(f"Errore nella creazione della tabella risultati: {str(e)}")
        error_run.font.size = Pt(7)
        error_run.font.name = 'Arial'
        error_run.font.color.rgb = RGBColor(255, 0, 0)  # Rosso per errore

def create_word_report(report_data, file_path):
    """Crea il documento Word con i dati del report in formato SEA SRLS"""
    
    try:
        debug_log(f"Inizio creazione documento Word: {file_path}")
        debug_log(f"Tipo report_data: {type(report_data)}")
        
        # Verifica cartella di destinazione
        dest_dir = os.path.dirname(file_path)
        if not os.path.exists(dest_dir):
            debug_log(f"Creazione cartella: {dest_dir}")
            os.makedirs(dest_dir, exist_ok=True)
        
        # Crea nuovo documento con codifica UTF-8
        doc = Document()
        debug_log("Documento Word creato")
        
        # FIX: Gestisci sia dizionari che liste
        if isinstance(report_data, dict):
            # Formato atteso: {'risultati_classificazione': [...], '_metadata': {...}}
            risultati = report_data.get('risultati_classificazione', [])
            metadata = report_data.get('_metadata', {})
            debug_log(f"Formato dizionario - Risultati: {len(risultati)}, Metadata: {list(metadata.keys())}")
        elif isinstance(report_data, list):
            # Formato diretto: [...]
            risultati = report_data
            metadata = {}
            debug_log(f"Formato lista - Risultati: {len(risultati)}")
        else:
            # Formato sconosciuto
            debug_log(f"Formato sconosciuto: {type(report_data)}")
            risultati = []
            metadata = {}
        
        # PULIZIA CARATTERI SPECIALI nei risultati
        for risultato in risultati:
            for key, value in risultato.items():
                if isinstance(value, str):
                    # Sostituisci caratteri problematici
                    value = value.replace('≥', '>=')
                    value = value.replace('•', ' - ')
                    value = value.replace('â‰¥', '>=')
                    value = value.replace('â€¢', ' - ')
                    risultato[key] = value
        
        # Estrai informazioni base
        codice_eer = 'N/A'
        numero_campionamento = 'N/A'
        data_campionamento = 'N/A'
        
        if isinstance(metadata.get('infoCertificato'), dict):
            codice_eer = metadata['infoCertificato'].get('codiceEER', 'N/A')
            numero_campionamento = metadata['infoCertificato'].get('numeroCampionamento', 'N/A')
            data_campionamento = metadata['infoCertificato'].get('dataCampionamento', 'N/A')
        
        debug_log(f"Info estratte - EER: {codice_eer}, Campionamento: {numero_campionamento}")
        
        # === COSTRUZIONE DOCUMENTO ===
        
        # 1. Aggiungi footer con data rapporto e numero pagina
        add_footer_with_page_numbers(doc, metadata)
        
        # 2. Header completo nel corpo del documento (NORMALE, NON GRIGIO)
        add_complete_header_content(doc, metadata)
        
        # 3. Sezione contenuto dopo la linea di separazione
        add_document_title_section(doc, metadata)
        
        # 4. Tabelle di classificazione
        create_classification_tables(doc, metadata, risultati)
        
        # 5. INTERRUZIONE DI PAGINA E HEADER RIPETUTO PER TABELLA 4
        debug_log("Aggiunta nuova pagina per Tabella 4 con header ripetuto")
        add_page_break_with_header(doc, metadata)
        
        # 6. Tabella risultati sulla nuova pagina
        create_results_table(doc, risultati)
        
        # 7. Footer informativo
        doc.add_page_break()
        
        footer_title = doc.add_paragraph()
        footer_title_run = footer_title.add_run('Informazioni Tecniche')
        footer_title_run.font.size = Pt(7)
        footer_title_run.font.bold = True
        footer_title_run.font.name = 'Arial'
        
        footer_info = doc.add_paragraph()
        footer_text = f"""Documento generato automaticamente dal sistema di classificazione rifiuti SEA SRLS.
Data e ora di generazione: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}

Per maggiori informazioni contattare:
SEA SRLS - Via Madre Teresa di Calcutta, 100 - 62029 Tolentino (MC)
Tel.: 0733/974947 - e-mail: info@seatoientino.it"""
        
        footer_run = footer_info.add_run(footer_text)
        footer_run.font.size = Pt(7)
        footer_run.font.name = 'Arial'
        
        debug_log("Footer completato")
        
        # Salva il documento
        debug_log(f"Salvataggio documento: {file_path}")
        doc.save(file_path)
        
        # Verifica che il file sia stato creato
        if os.path.exists(file_path):
            file_size = os.path.getsize(file_path)
            debug_log(f"File creato con successo, dimensione: {file_size} bytes")
            return True
        else:
            debug_log("ERRORE: File non creato")
            return False
        
    except Exception as e:
        debug_log(f"Errore durante creazione documento: {e}")
        import traceback
        debug_log(f"Traceback: {traceback.format_exc()}")
        return False

def main():
    """Funzione principale"""
    try:
        debug_log("Inizio main()")
        
        # Leggi dati da stdin
        debug_log("Lettura dati da stdin...")
        input_data = sys.stdin.read()
        debug_log(f"Dati ricevuti: {len(input_data)} caratteri")
        
        if not input_data.strip():
            debug_log("ERRORE: Nessun dato ricevuto da stdin")
            print("ERROR: Nessun dato ricevuto")
            sys.exit(1)
        
        debug_log("Parsing JSON...")
        data = json.loads(input_data)
        
        report_data = data['reportData']
        file_path = data['filePath']
        
        debug_log(f"Report data keys: {list(report_data.keys()) if isinstance(report_data, dict) else 'Not a dict'}")
        debug_log(f"File path: {file_path}")
        
        # Crea il report
        debug_log("Avvio creazione report...")
        success = create_word_report(report_data, file_path)
        
        if success:
            print("SUCCESS: Report Word creato con successo")
            debug_log("SUCCESS: Report completato")
        else:
            print("ERROR: Errore nella creazione del report")
            debug_log("ERROR: Creazione fallita")
            sys.exit(1)
            
    except json.JSONDecodeError as e:
        debug_log(f"Errore JSON: {e}")
        print(f"ERROR: Errore parsing JSON: {e}")
        sys.exit(1)
    except KeyError as e:
        debug_log(f"Errore chiave mancante: {e}")
        print(f"ERROR: Chiave mancante nei dati: {e}")
        sys.exit(1)
    except Exception as e:
        debug_log(f"Errore generale: {e}")
        import traceback
        debug_log(f"Traceback completo: {traceback.format_exc()}")
        print(f"ERROR: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()